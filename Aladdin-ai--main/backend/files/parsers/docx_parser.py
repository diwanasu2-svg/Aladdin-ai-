"""DOCX parser — text, tables, metadata."""
from __future__ import annotations
import asyncio, io, logging
from typing import Any, Dict
log = logging.getLogger(__name__)

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False
    log.warning("python-docx not installed")


class DOCXParser:
    @property
    def available(self): return _DOCX

    async def parse(self, data: bytes) -> Dict[str, Any]:
        if not _DOCX:
            return {"error": "python-docx not installed", "text": ""}
        def _run():
            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                tables.append({"rows": rows})
            core = doc.core_properties
            return {
                "text": "\n".join(paragraphs),
                "paragraphs": len(paragraphs),
                "tables": tables,
                "metadata": {"author": getattr(core, "author", ""),
                              "title": getattr(core, "title", ""),
                              "created": str(getattr(core, "created", ""))},
            }
        return await asyncio.get_running_loop().run_in_executor(None, _run)
